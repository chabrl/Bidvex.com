"""
iter319 — Global onboarding + Claude auto-screening tests.

Focused on pure-function logic that doesn't need a live Claude call.
The integration test for the actual Claude round-trip is in the live
HTTP suite (runs against /api/admin/careers/applicants/{id}/screen).
"""
import json
import sys
from pathlib import Path

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

import pytest

from services import careers_screening as cs


# ─── _parse_screening_json ─────────────────────────────────────────────

class TestParseScreeningJson:
    def test_plain_json_object(self):
        raw = '{"summary":"Solid candidate","recommendation":"Yes","key_signals":["A","B"]}'
        out = cs._parse_screening_json(raw)
        assert out["status"] == "ok"
        assert out["summary"] == "Solid candidate"
        assert out["recommendation"] == "Yes"
        assert out["key_signals"] == ["A", "B"]

    def test_json_wrapped_in_markdown_fences(self):
        raw = '```json\n{"summary":"X","recommendation":"Maybe"}\n```'
        out = cs._parse_screening_json(raw)
        assert out["status"] == "ok"
        assert out["recommendation"] == "Maybe"

    def test_json_with_surrounding_prose(self):
        raw = 'Here is my analysis:\n{"summary":"Y","recommendation":"no"}\nDone.'
        out = cs._parse_screening_json(raw)
        assert out["status"] == "ok"
        assert out["recommendation"] == "No"  # case-normalised

    def test_invalid_recommendation_defaults_to_maybe(self):
        raw = '{"summary":"Z","recommendation":"strong yes"}'
        out = cs._parse_screening_json(raw)
        assert out["recommendation"] == "Maybe"

    def test_malformed_json_returns_failed_envelope(self):
        out = cs._parse_screening_json("totally not json")
        assert out["status"] == "failed"
        assert "json parse" in out["error"]

    def test_summary_truncated_to_220_chars(self):
        long_summary = "x" * 500
        raw = json.dumps({"summary": long_summary, "recommendation": "Yes"})
        out = cs._parse_screening_json(raw)
        assert len(out["summary"]) <= 220


# ─── _clean_text ───────────────────────────────────────────────────────

class TestCleanText:
    def test_collapses_whitespace(self):
        out = cs._clean_text("foo   bar\n\n\n\nbaz")
        assert out == "foo bar\n\nbaz"

    def test_caps_length(self):
        big = "x" * (cs.RESUME_TEXT_MAX_CHARS + 1000)
        out = cs._clean_text(big)
        assert len(out) <= cs.RESUME_TEXT_MAX_CHARS + 50  # + truncation suffix


# ─── extract_text_from_file (real PDF + DOCX round-trip) ───────────────

def _make_test_pdf(tmp_path: Path) -> Path:
    """Build a real PDF with reportlab so pypdf can extract text from it."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed")
    p = tmp_path / "cv.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.setFont("Helvetica", 11)
    t = c.beginText(40, 750)
    for line in [
        "Test Candidate",
        "8 years outbound BDR experience.",
        "Bilingual French and English.",
        "Salesforce, HubSpot, Aircall dialer.",
    ]:
        t.textLine(line)
    c.drawText(t)
    c.save()
    return p


def _make_test_docx(tmp_path: Path) -> Path:
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")
    p = tmp_path / "cv.docx"
    d = docx.Document()
    d.add_paragraph("Test Candidate")
    d.add_paragraph("8 years outbound BDR experience.")
    d.add_paragraph("Bilingual French and English.")
    d.save(str(p))
    return p


class TestExtractTextFromFile:
    def test_pdf_extraction(self, tmp_path):
        p = _make_test_pdf(tmp_path)
        out = cs.extract_text_from_file(p)
        assert "Test Candidate" in out
        assert "BDR" in out

    def test_docx_extraction(self, tmp_path):
        p = _make_test_docx(tmp_path)
        out = cs.extract_text_from_file(p)
        assert "Test Candidate" in out
        assert "Bilingual" in out

    def test_unsupported_ext_returns_empty(self, tmp_path):
        p = tmp_path / "f.bin"
        p.write_bytes(b"\x00\x01\x02")
        assert cs.extract_text_from_file(p) == ""


# ─── Global onboarding validation logic (mirror of route handler) ──────

def _emulate_country_validation(*, country, province, state):
    """Mirror of the apply route's location validation rules. Returns
    the error code string (or '' on success)."""
    country = (country or "").strip()
    if not country:
        return "missing_country"
    if country == "Canada" and not (province or "").strip():
        return "missing_province"
    if country == "United States" and not (state or "").strip():
        return "missing_state"
    return ""


class TestGlobalOnboardingValidation:
    def test_canada_requires_province(self):
        assert _emulate_country_validation(country="Canada", province="", state="") == "missing_province"
        assert _emulate_country_validation(country="Canada", province="QC", state="") == ""

    def test_us_requires_state(self):
        assert _emulate_country_validation(country="United States", province="", state="") == "missing_state"
        assert _emulate_country_validation(country="United States", province="", state="TX") == ""

    def test_other_country_no_secondary_required(self):
        # France, UK, Singapore — none require a secondary tier.
        for c in ("France", "United Kingdom", "Singapore", "Japan", "Lebanon"):
            assert _emulate_country_validation(country=c, province="", state="") == ""

    def test_missing_country_blocks(self):
        assert _emulate_country_validation(country="", province="QC", state="") == "missing_country"


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
